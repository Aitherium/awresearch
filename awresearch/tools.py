"""Curated tool set for the Deep Research Agent.

These are the ONLY tools the agent may call (default-deny; the brain pack lists
them as a whitelist). Read & generate only — no shell, no file delete, no deploy.

Each tool is a closure over a ResearchSession so it can:
  - reach the knowledge graph (adk GraphMemory, SQLite, offline-safe),
  - feed the SavingsLedger (so the UI can show tokens used vs. saved),
  - share fetched-page and findings caches across the session (dedup).

`build_research_tools(session)` returns a list of plain functions; serve.py
registers them on the AitherAgent via agent._tools.register(fn).
"""

from __future__ import annotations

import asyncio
import json
import logging
import re
import time
from collections import Counter
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any, Callable

import httpx

from . import aithersearch, sources, trust
from .ledger import SavingsLedger

logger = logging.getLogger("deep_research.tools")


@dataclass
class ResearchSession:
    """Shared state for one research session: graph, ledger, caches, artifacts."""

    graph: Any                       # adk.graph_memory.GraphMemory
    ledger: SavingsLedger
    artifacts_dir: Path
    session_id: str = "session"
    _page_cache: dict[str, str] = field(default_factory=dict)
    _finding_keys: set[str] = field(default_factory=set)
    sources: list[dict] = field(default_factory=list)  # ordered, for citations
    contradictions: list[dict] = field(default_factory=list)  # cross-source conflicts
    _contra_keys: set[str] = field(default_factory=set)

    def cite(self, title: str, url: str) -> int:
        """Register a source and return its 1-based citation number.

        Stamps a cheap url-only authority tier on first sight (tier/authority/
        domain) so the UI gets source trust for free; freshness fills in later
        once the page is actually read (note_source_text)."""
        for i, s in enumerate(self.sources):
            if s["url"] == url:
                return i + 1
        s = {"title": title, "url": url}
        try:
            a = trust.domain_authority(url)
            s.update(domain=a["domain"], tier=a["tier"], authority=a["score"])
        except Exception:  # noqa: BLE001
            pass
        self.sources.append(s)
        return len(self.sources)

    def note_source_text(self, url: str, text: str) -> None:
        """Once a registered source's page is read, fill in its freshness/trust
        from the body text + URL (date decay). No-op if the URL isn't a source;
        never raises."""
        try:
            for s in self.sources:
                if s.get("url") == url:
                    f = trust.trust_score(url, (text or "")[:4000])
                    s["date"] = f["date"]
                    s["age_days"] = f["age_days"]
                    s["freshness"] = f["freshness"]
                    s["trust"] = f["score"]
                    s.setdefault("authority", f["authority"])
                    s.setdefault("tier", f["tier"])
                    s.setdefault("domain", f["domain"])
                    return
        except Exception:  # noqa: BLE001
            pass


def _strip_html(html: str) -> str:
    html = re.sub(r"<(script|style)[^>]*>.*?</\1>", " ", html, flags=re.DOTALL | re.IGNORECASE)
    html = re.sub(r"<[^>]+>", " ", html)
    html = re.sub(r"\s+", " ", html)
    return html.strip()


_BOILERPLATE = ("script", "style", "nav", "header", "footer", "aside", "form",
                "noscript", "iframe", "svg", "button")


def _extract_main(html: str) -> str:
    """Readability-style main-text extraction with lxml: drop boilerplate (nav/footer/
    ads/scripts), prefer the <article>/<main> or densest <body> block, and emit clean
    text that PRESERVES paragraph breaks, list items and table rows — far better signal
    for the LLM than a flat tag-strip. Falls back to _strip_html on any failure."""
    try:
        from lxml import html as lh
        doc = lh.fromstring(html)
        for el in doc.xpath("//*[%s]" % " or ".join(f"self::{t}" for t in _BOILERPLATE)):
            el.getparent().remove(el) if el.getparent() is not None else None
        # also drop obvious chrome by class/id
        for el in doc.xpath('//*[contains(translate(@class,"ABCDEFGHIJKLMNOPQRSTUVWXYZ","abcdefghijklmnopqrstuvwxyz"),"cookie") '
                            'or contains(translate(@id,"ABCDEFGHIJKLMNOPQRSTUVWXYZ","abcdefghijklmnopqrstuvwxyz"),"sidebar") '
                            'or contains(translate(@class,"ABCDEFGHIJKLMNOPQRSTUVWXYZ","abcdefghijklmnopqrstuvwxyz"),"advert")]'):
            if el.getparent() is not None:
                el.getparent().remove(el)
        # choose the richest container
        roots = doc.xpath("//article") or doc.xpath("//main") or doc.xpath("//body") or [doc]
        root = max(roots, key=lambda r: len(r.text_content() or ""))
        parts: list[str] = []
        for node in root.xpath(".//h1|.//h2|.//h3|.//h4|.//p|.//li|.//tr|.//pre|.//blockquote|.//figcaption"):
            tag = node.tag
            if tag == "tr":
                cells = [re.sub(r"\s+", " ", c.text_content()).strip()
                         for c in node.xpath("./td|./th")]
                line = " | ".join(c for c in cells if c)
            else:
                line = re.sub(r"\s+", " ", node.text_content() or "").strip()
            if not line:
                continue
            if tag in ("h1", "h2", "h3", "h4"):
                line = "\n## " + line
            elif tag == "li":
                line = "- " + line
            parts.append(line)
        text = "\n".join(parts).strip()
        # if structure-aware extraction got too little (JS-only page, odd markup), fall back
        if len(text) < 200:
            flat = _strip_html(html)
            return flat if len(flat) > len(text) else text
        return re.sub(r"\n{3,}", "\n\n", text)
    except Exception:  # noqa: BLE001
        return _strip_html(html)


def _as_int(v, default: int) -> int:
    """Coerce an LLM-supplied arg to int. Models sometimes pass lists/strings/None
    (e.g. `angles=["overview","recent"]`); never crash on that."""
    if isinstance(v, bool):
        return default
    if isinstance(v, int):
        return v
    if isinstance(v, float):
        return int(v)
    if isinstance(v, (list, tuple, set)):
        return len(v) or default
    if isinstance(v, str):
        m = re.search(r"-?\d+", v)
        return int(m.group()) if m else default
    return default


# ── semantic dedup / contradiction heuristics (cheap, offline, no LLM) ────────
_STOP = {
    "the", "a", "an", "and", "or", "of", "to", "in", "on", "for", "with", "by",
    "at", "as", "is", "are", "was", "were", "be", "been", "it", "its", "this",
    "that", "these", "those", "from", "has", "have", "had", "but", "than",
    "into", "about", "over", "their", "they", "them", "which", "while", "also",
}
# negation / refutation cues — a flip between two related claims signals conflict.
_NEG = {
    "not", "no", "never", "cannot", "cant", "wont", "isnt", "arent", "doesnt",
    "didnt", "none", "without", "false", "fails", "failed", "denies", "denied",
    "rejected", "unproven", "refutes", "contradicts", "debunked",
}

DEDUP_SIM = 0.82      # token-Jaccard at/above this = a near-duplicate fact
CONFLICT_SIM = 0.70   # only claims THIS similar are the same assertion -> safe to test for conflict
RELATE_SIM = 0.45     # (kept for reference; below CONFLICT_SIM is "same topic, different claim" -> ignore)

_NUM_RE = re.compile(
    r"\b\d[\d,]*(?:\.\d+)?\s?(?:%|percent|million|billion|bn|m|k|x|gwh|kwh|wh|mw|gw)?\b")


def _norm_tokens(s: str) -> set[str]:
    """Content tokens for Jaccard: alnum/%$. tokens, stopwords dropped, len>=3
    (pure-number / % tokens kept regardless of length)."""
    out: set[str] = set()
    for t in re.findall(r"[a-z0-9%$.]+", (s or "").lower()):
        t = t.strip(".")
        if not t or t in _STOP:
            continue
        if len(t) >= 3 or any(ch.isdigit() for ch in t) or t.endswith("%"):
            out.add(t)
    return out


def _has_neg(s: str) -> bool:
    toks = {t.replace("'", "") for t in re.findall(r"[a-z']+", (s or "").lower())}
    return bool(toks & _NEG)


def _numbers(s: str) -> set[str]:
    """Normalized numeric/magnitude tokens (strip commas/spaces) for conflict test."""
    out: set[str] = set()
    for m in _NUM_RE.findall((s or "").lower()):
        n = m.replace(",", "").replace(" ", "").strip()
        if n and any(ch.isdigit() for ch in n):
            out.add(n)
    return out


def _jaccard(a: set, b: set) -> float:
    if not a or not b:
        return 0.0
    return len(a & b) / len(a | b)


def _strip_source(c: str) -> str:
    """Drop a trailing ' [source: …]' tag from a stored finding's content."""
    return re.sub(r"\s*\[source:[^\]]*\]\s*$", "", c or "").strip()


def build_research_tools(session: ResearchSession) -> list[Callable]:
    graph = session.graph
    ledger = session.ledger

    async def web_search(query: str, limit: int = 6) -> str:
        """Search the open web (DuckDuckGo, no API key). Returns titles, URLs, and snippets.

        query: what to search for
        limit: max results (default 6)
        """
        limit = _as_int(limit, 6)
        ledger.note_search()
        res = await aithersearch.search(str(query), limit=limit)
        # Register results as citable sources up front.
        for r in res.get("results", []):
            if r.get("url"):
                session.cite(r.get("title", ""), r["url"])
        return json.dumps(res)

    async def deep_research(topic: str, angles: int = 3) -> str:
        """Run a multi-angle web sweep on a topic and return de-duplicated results.

        Fires several complementary queries (overview, recent developments,
        comparisons/criticism) and merges the hits. Use for broad questions.

        topic: the subject to investigate
        angles: how many query angles to run (default 3, max 5)
        """
        angles = max(1, min(_as_int(angles, 3), 5))
        topic = str(topic)
        templates = [
            "{t}",
            "{t} latest developments 2026",
            "{t} comparison OR analysis",
            "{t} criticism OR limitations OR risks",
            "{t} data OR statistics OR report",
        ][:angles]
        merged: dict[str, dict] = {}
        for tmpl in templates:
            ledger.note_search()
            res = await aithersearch.search(tmpl.format(t=topic), limit=5)
            for r in res.get("results", []):
                u = r.get("url", "")
                if u and u not in merged:
                    merged[u] = r
                    session.cite(r.get("title", ""), u)
        return json.dumps({"topic": topic, "queries": len(templates),
                           "results": list(merged.values()), "count": len(merged)})

    async def _read_page(url: str, client: httpx.AsyncClient) -> dict:
        """Fetch + extract one page (shared by fetch_url and fetch_many). Returns a dict
        with text|error; populates the per-session cache and the savings ledger."""
        url = str(url)
        if url in session._page_cache:
            return {"url": url, "text": session._page_cache[url], "cached": True}
        try:
            resp = await client.get(url, headers={"User-Agent":
                "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 "
                "(KHTML, like Gecko) Chrome/122.0 Safari/537.36"})
            resp.raise_for_status()
            ctype = resp.headers.get("content-type", "")
            text = resp.text if "html" in ctype or not ctype else resp.text
            text = _extract_main(text)
        except Exception as exc:  # noqa: BLE001
            return {"url": url, "error": str(exc)}
        session._page_cache[url] = text
        ledger.note_page_read()
        try:
            session.note_source_text(url, text)   # fill freshness/trust now that we read it
        except Exception:  # noqa: BLE001
            pass
        return {"url": url, "text": text, "cached": False, "length": len(text)}

    async def fetch_url(url: str, max_chars: int = 6000) -> str:
        """Fetch a web page and return its readable MAIN text (boilerplate removed;
        headings, lists and tables preserved).

        Reuses a per-session cache: re-fetching the same URL costs nothing and is
        credited as tokens saved.

        url: the page to read
        max_chars: max characters of text to return (default 6000)
        """
        max_chars = _as_int(max_chars, 6000)
        async with httpx.AsyncClient(timeout=20.0, follow_redirects=True, max_redirects=5) as client:
            doc = await _read_page(url, client)
        if doc.get("cached"):
            ledger.record_dedup(doc["text"][:max_chars])      # avoided a re-fetch + re-read
        if "error" in doc:
            return json.dumps(doc)
        return json.dumps({**doc, "text": doc["text"][:max_chars]})

    async def fetch_many(urls, max_chars: int = 5000) -> str:
        """Read SEVERAL pages CONCURRENTLY in one call — the fast way to gather a topic.
        Pass the best 3-8 result URLs and get all their main texts back at once instead
        of fetching one at a time.

        urls: a list of URLs (or a comma/newline-separated string)
        max_chars: max characters per page (default 5000)
        """
        max_chars = _as_int(max_chars, 5000)
        if isinstance(urls, str):
            urls = [u.strip() for u in re.split(r"[,\n]+", urls) if u.strip()]
        urls = [str(u) for u in (urls or [])][:8]              # bound the fan-out
        if not urls:
            return json.dumps({"pages": [], "count": 0})
        async with httpx.AsyncClient(timeout=20.0, follow_redirects=True, max_redirects=5) as client:
            docs = await asyncio.gather(*[_read_page(u, client) for u in urls],
                                        return_exceptions=True)
        pages = []
        for u, d in zip(urls, docs):
            if isinstance(d, Exception):
                pages.append({"url": u, "error": str(d)})
            elif d.get("cached"):
                ledger.record_dedup(d["text"][:max_chars])
                pages.append({"url": u, "text": d["text"][:max_chars], "cached": True})
            elif "error" in d:
                pages.append(d)
            else:
                pages.append({"url": u, "text": d["text"][:max_chars], "length": d.get("length")})
        return json.dumps({"pages": pages, "count": len(pages)})

    async def save_finding(claim: str, source_url: str = "", topic: str = "") -> str:
        """Store a verified fact + its source in the knowledge graph for reuse.

        Call this for every solid fact you find so you never re-research it.

        claim: the fact, in one sentence
        source_url: the URL it came from
        topic: optional grouping label
        """
        # (a) exact-key fast path FIRST (free — no embed, no graph hit).
        key = re.sub(r"\s+", " ", claim.strip().lower())[:160]
        if key in session._finding_keys:
            ledger.record_dedup(claim)  # already known — no re-store/re-embed
            return json.dumps({"stored": False, "duplicate": True})
        session._finding_keys.add(key)

        # (c) SEMANTIC PASS — graph.search is embedding-backed (offline fallback
        # to feature-hash when Ollama is down); confirm with our own token-Jaccard
        # over the candidate shortlist. Fully offline-safe + wrapped.
        dup = None
        contra = None
        try:
            cands = await graph.search(claim, limit=4)
            claim_tok = _norm_tokens(claim)
            claim_frame = {t for t in claim_tok if not t.isdigit()}   # non-numeric "assertion frame"
            claim_num = _numbers(claim)
            claim_neg = _has_neg(claim)
            for n in cands:
                ctext = _strip_source(n.content or "")
                ctok = _norm_tokens(ctext)
                # Conflict/dedup only between the SAME ASSERTION — measured on the
                # non-numeric FRAME so "shipped 200k cells" vs "shipped 50k cells"
                # (frame identical, numbers differ -> a real contradiction) still
                # matches, while two DIFFERENT facts that merely share a topic word
                # and each carry a number (low frame overlap) are NOT false-flagged.
                frame = _jaccard(claim_frame, {t for t in ctok if not t.isdigit()})
                if frame < CONFLICT_SIM:
                    continue
                # A polarity flip (negation) or a disagreeing quantity on that same
                # assertion is a CONTRADICTION — tested before dedup so a merely
                # negated near-identical sentence isn't swallowed as a duplicate.
                neg = claim_neg != _has_neg(ctext)
                cn = _numbers(ctext)
                num = bool((claim_num - cn) and (cn - claim_num))  # each has a number the other lacks
                if neg or num:
                    if contra is None:
                        contra = (n, "negation" if neg else "number")
                    continue
                if _jaccard(claim_tok, ctok) >= DEDUP_SIM:   # identical incl. numbers -> duplicate
                    dup = n
                    break
        except Exception as exc:  # noqa: BLE001
            logger.debug("save_finding semantic pass failed: %s", exc)

        # (d) near-duplicate: honestly avoided a re-store/re-embed of the same fact.
        if dup is not None:
            ledger.record_dedup(claim)
            return json.dumps({"stored": False, "duplicate": True,
                               "semantic": True, "matched": dup.label})

        # (e) store the new fact.
        tags = [t for t in ["finding", topic] if t]
        new = None
        try:
            new = await graph.add_node(
                label=(topic or claim[:60]),
                node_type="fact",
                content=f"{claim} [source: {source_url}]" if source_url else claim,
                tags=tags,
                source_session=session.session_id,
            )
        except Exception as exc:  # noqa: BLE001
            logger.debug("graph.add_node failed: %s", exc)
        ledger.note_finding()
        cite_n = session.cite("", source_url) if source_url else 0
        out = {"stored": True, "citation": cite_n}

        # record a contradiction (a NEW node — credits no savings) so the model
        # SEES the disagreement inline this turn and can surface it before FINAL.
        if contra is not None:
            other = contra[0]
            ck = "||".join(sorted([claim[:80], (other.content or "")[:80]]))
            if ck not in session._contra_keys:
                session._contra_keys.add(ck)
                other_text = _strip_source(other.content or "")
                src_b = ""
                msrc = re.search(r"\[source:\s*([^\]]*)\]", other.content or "")
                if msrc:
                    src_b = msrc.group(1).strip()
                conflict = {"claim_a": claim, "claim_b": other_text,
                            "kind": contra[1], "source_a": source_url,
                            "source_b": src_b}
                session.contradictions.append(conflict)
                if new is not None:
                    try:
                        await graph.add_edge(new.id, other.id, "contradicts",
                                             metadata={"kind": contra[1]})
                    except Exception as exc:  # noqa: BLE001
                        logger.debug("contradiction edge failed: %s", exc)
                out["contradiction"] = conflict
        return json.dumps(out)

    async def recall(query: str, limit: int = 5) -> str:
        """Recall what you already learned this session from memory + the graph.

        ALWAYS call this before searching — if you already know it, reuse it and
        cite the cached finding instead of spending another search. Reused tokens
        are counted as 'saved by memory'.

        query: what you're trying to remember
        limit: max results (default 5)
        """
        limit = _as_int(limit, 5)
        try:
            nodes = await graph.search(str(query), limit=limit)
        except Exception as exc:  # noqa: BLE001
            return json.dumps({"results": [], "error": str(exc)})
        results = []
        reused_text = ""
        for n in nodes:
            item = {"label": n.label, "content": n.content, "type": n.node_type}
            results.append(item)
            reused_text += " " + (n.content or "")
        if results:
            ledger.record_recall(reused_text)  # answered from memory, not a fresh fetch
        return json.dumps({"results": results, "count": len(results),
                           "reused_from_memory": bool(results)})

    async def knowledge_graph() -> str:
        """Show what the agent has learned this session (graph stats + counts),
        including any cross-source contradictions found and a source-trust roll-up."""
        try:
            stats = await graph.get_stats()
        except Exception as exc:  # noqa: BLE001
            stats = {"error": str(exc)}
        by_tier = Counter(s.get("tier", "unknown") for s in session.sources)
        top = []
        for i, s in enumerate(session.sources[:8]):
            row = {"n": i + 1, "domain": s.get("domain", ""),
                   "tier": s.get("tier", "unknown"),
                   "authority": s.get("authority"), "url": s.get("url", "")}
            if s.get("freshness") is not None:
                row["freshness"] = s.get("freshness")
            if s.get("trust") is not None:
                row["trust"] = s.get("trust")
            top.append(row)
        return json.dumps({
            "graph": stats,
            "sources_cited": len(session.sources),
            "contradictions": len(session.contradictions),
            "contradiction_list": session.contradictions[:10],
            "sources_by_tier": dict(by_tier),
            "top_sources": top,
        })

    async def contradictions() -> str:
        """Surface the conflicts found between sources this session.

        Returns the running conflict ledger (each: claim_a, claim_b, kind,
        source_a, source_b) so you can state 'sources disagree' EXPLICITLY,
        with both sides cited, before writing FINAL. No args, no network.
        """
        return json.dumps({"contradictions": session.contradictions,
                           "count": len(session.contradictions)})

    # ── Specialized & recovery sources (keyless: papers + archive recovery) ──

    async def scholarly_search(query: str, limit: int = 8) -> str:
        """Search the SCHOLARLY literature (arXiv + Semantic Scholar + Crossref,
        no API key). Returns peer-reviewed / preprint papers with title, authors,
        year, venue, abstract, and DOI/URL. Use for academic, scientific, or
        technical claims that need primary literature.

        query: what to search for
        limit: max papers (default 8)
        """
        limit = _as_int(limit, 8)
        res = await sources.scholarly_search(str(query), limit=limit)
        # honest: one search credit per provider that ACTUALLY answered — and none at
        # all if every provider failed (don't credit work that produced nothing).
        for _ in (res.get("providers") or []):
            ledger.note_search()
        for r in res.get("papers", []):
            if r.get("url"):
                session.cite(r.get("title", ""), r["url"])
        return json.dumps(res)

    async def wayback_lookup(url: str, limit: int = 10) -> str:
        """Recover an archived snapshot of a DEAD or moved URL via the Internet
        Archive Wayback Machine (no API key). Returns the closest live snapshot
        plus a timeline of captures — time-travel to primary sources that vanished
        or changed. The snapshot URLs are fetchable by fetch_url / fetch_many.

        url: the dead/moved URL to recover
        limit: max captures in the timeline (default 10)
        """
        limit = _as_int(limit, 10)
        res = await sources.wayback_lookup(str(url), limit=limit)
        closest = res.get("closest") or {}
        caps = res.get("captures") or []
        if closest.get("snapshot_url") or caps:
            ledger.note_search()        # credit only a lookup that actually recovered something
        if closest.get("snapshot_url"):
            session.cite(f"Wayback: {url}", closest["snapshot_url"])
        if caps and caps[0].get("snapshot_url"):
            session.cite(f"Wayback capture: {url}", caps[0]["snapshot_url"])
        return json.dumps(res)

    async def filings_search(query: str, limit: int = 10) -> str:
        """Search SEC EDGAR public-company filings (10-K / 10-Q / 8-K / S-1 / prospectuses,
        no API key) by full text. Returns filings with company, form type, date, and a
        direct document URL (read it with fetch_url). Use for primary-source corporate,
        financial, and regulatory facts.

        query: what to search the filings for
        limit: max filings (default 10)
        """
        limit = _as_int(limit, 10)
        res = await sources.filings_search(str(query), limit=limit)
        if res.get("filings"):
            ledger.note_search()
        for f in res.get("filings", []):
            if f.get("url"):
                session.cite(f"{f.get('company', '')} {f.get('form', '')} {f.get('date', '')}".strip(),
                             f["url"])
        return json.dumps(res)

    async def code_search(query: str, limit: int = 10) -> str:
        """Search GitHub repositories (no API key, sorted by stars) for implementations,
        tools, and open-source state of the art. Returns repos with name, description,
        stars, language, URL and last-push date.

        query: what to search for (GitHub query syntax allowed, e.g. 'rag language:python')
        limit: max repos (default 10)
        """
        limit = _as_int(limit, 10)
        res = await sources.code_search(str(query), limit=limit)
        if res.get("repos"):
            ledger.note_search()
        for r in res.get("repos", []):
            if r.get("url"):
                session.cite(r.get("name", ""), r["url"])
        return json.dumps(res)

    async def wikidata_lookup(query: str, limit: int = 8) -> str:
        """Resolve a name/term to canonical Wikidata entities (stable Q-IDs, labels,
        descriptions). Use to DISAMBIGUATE and ground an entity (person, company, place,
        concept) before researching it, so you search the right thing.

        query: the name/term to resolve
        limit: max entities (default 8)
        """
        limit = _as_int(limit, 8)
        res = await sources.wikidata_lookup(str(query), limit=limit)
        if res.get("entities"):
            ledger.note_search()
        return json.dumps(res)

    # ── Report generation (read/generate only; writes to the artifacts dir) ──

    def _sources_block() -> str:
        if not session.sources:
            return ""
        lines = ["", "## Sources", ""]
        for i, s in enumerate(session.sources, 1):
            title = s.get("title") or s.get("url")
            lines.append(f"{i}. {title} — {s.get('url')}")
        return "\n".join(lines)

    def _save(name: str, data: bytes) -> Path:
        session.artifacts_dir.mkdir(parents=True, exist_ok=True)
        ts = time.strftime("%Y%m%d-%H%M%S")
        safe = re.sub(r"[^a-zA-Z0-9_.-]+", "-", name).strip("-") or "report"
        path = session.artifacts_dir / f"{ts}-{safe}"
        path.write_bytes(data)
        return path

    def generate_markdown(title: str, content: str) -> str:
        """Write a Markdown research report (auto-appends a Sources section).

        title: report title
        content: report body in Markdown (use [1], [2] inline citations)
        """
        md = f"# {title}\n\n{content}\n{_sources_block()}\n"
        path = _save(f"{title}.md", md.encode("utf-8"))
        return json.dumps({"saved": True, "format": "markdown",
                           "path": str(path), "file": path.name})

    def generate_pdf(title: str, content: str) -> str:
        """Write a PDF research report (reportlab). Auto-appends Sources.

        title: report title
        content: report body in Markdown-ish text (#, ##, - bullets, [1] citations)
        """
        try:
            from reportlab.lib.pagesizes import letter
            from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
            from reportlab.lib.units import inch
            from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer
        except ImportError:
            return json.dumps({"saved": False, "error": "reportlab not installed"})

        import io
        buf = io.BytesIO()
        doc = SimpleDocTemplate(buf, pagesize=letter, title=title,
                                topMargin=0.8 * inch, bottomMargin=0.8 * inch)
        styles = getSampleStyleSheet()
        h1 = ParagraphStyle("H1c", parent=styles["Heading1"], spaceAfter=12)
        story = [Paragraph(title, h1), Spacer(1, 8)]
        body = content + "\n" + _sources_block()
        for raw in body.split("\n"):
            line = raw.rstrip()
            if not line:
                story.append(Spacer(1, 6))
            elif line.startswith("## "):
                story.append(Paragraph(line[3:], styles["Heading2"]))
            elif line.startswith("# "):
                story.append(Paragraph(line[2:], styles["Heading1"]))
            elif line.startswith("- "):
                story.append(Paragraph(f"&bull;&nbsp;{_esc(line[2:])}", styles["BodyText"]))
            else:
                story.append(Paragraph(_esc(line), styles["BodyText"]))
        doc.build(story)
        path = _save(f"{title}.pdf", buf.getvalue())
        return json.dumps({"saved": True, "format": "pdf",
                           "path": str(path), "file": path.name})

    def generate_docx(title: str, content: str) -> str:
        """Write a Word (.docx) research report (python-docx). Auto-appends Sources.

        title: report title
        content: report body (#, ##, - bullets, [1] citations)
        """
        try:
            from docx import Document
        except ImportError:
            return json.dumps({"saved": False, "error": "python-docx not installed"})
        import io
        document = Document()
        document.add_heading(title, level=0)
        body = content + "\n" + _sources_block()
        for raw in body.split("\n"):
            line = raw.rstrip()
            if not line:
                continue
            if line.startswith("## "):
                document.add_heading(line[3:], level=2)
            elif line.startswith("# "):
                document.add_heading(line[2:], level=1)
            elif line.startswith("- "):
                document.add_paragraph(line[2:], style="List Bullet")
            else:
                document.add_paragraph(line)
        buf = io.BytesIO()
        document.save(buf)
        path = _save(f"{title}.docx", buf.getvalue())
        return json.dumps({"saved": True, "format": "docx",
                           "path": str(path), "file": path.name})

    return [
        web_search, deep_research, fetch_url, fetch_many, save_finding, recall,
        knowledge_graph, contradictions, scholarly_search, wayback_lookup,
        filings_search, code_search, wikidata_lookup,
        generate_markdown, generate_pdf, generate_docx,
    ]


def _esc(text: str) -> str:
    """Escape the few chars reportlab's mini-markup cares about."""
    return (text or "").replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
